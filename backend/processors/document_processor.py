# -*- coding: utf-8 -*-
"""
v4.7 — استخراج نص صافٍ خاطف: قراءة ميكانيكية فقط، بلا LLM ولا فلترة دلالية.
"""

import io
import json
import os
import re
from typing import Callable, List, Optional, Tuple

import logging

from processors.pdf_arabic_cleanup import (
    arabic_char_count,
    clean_pdf_arabic_text_with_fallback,
    pick_best_raw_pdf_text,
)

logger = logging.getLogger(__name__)

try:
    import orjson
except ImportError:
    orjson = None  # type: ignore[misc, assignment]

# صيغ مدعومة لرفع المستندات
ALLOWED_EXTENSIONS = {
    ".docx",
    ".pdf",
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".rtf",
}

EXTENSION_LABELS = {
    ".docx": "Word",
    ".pdf": "PDF",
    ".txt": "نص",
    ".md": "Markdown",
    ".markdown": "Markdown",
    ".json": "JSON",
    ".rtf": "RTF",
}


def _loads_json(file_bytes: bytes):
    """Parse JSON — orjson on bytes when available."""
    if orjson is not None:
        return orjson.loads(file_bytes)
    return json.loads(file_bytes.decode("utf-8-sig"))


def _json_sections_to_text(payload: dict) -> Optional[str]:
    """
    استخراج فوري من بنية Docx_to_JSON — منطق خفيف معزول عن json_document_parser
    (لا يستورد summarizer ولا clustering ولا audit).
    """
    content = payload.get("content")
    if not isinstance(content, dict):
        return None

    sections = content.get("sections")
    if not isinstance(sections, list) or not sections:
        return None

    parts: list[str] = []
    for section in sections:
        if not isinstance(section, dict):
            continue
        text = section.get("text")
        if isinstance(text, str) and text.strip():
            parts.append(text.strip())
        for item in section.get("items") or []:
            if item is not None:
                item_text = str(item).strip()
                if item_text:
                    parts.append(item_text)

    return "\n\n".join(parts) if parts else None


class DocumentProcessor:
    ALLOWED_EXTENSIONS = ALLOWED_EXTENSIONS

    @staticmethod
    def get_extension(filename: str) -> str:
        return os.path.splitext(filename or "")[1].lower()

    @staticmethod
    def is_allowed(filename: str) -> bool:
        return DocumentProcessor.get_extension(filename) in ALLOWED_EXTENSIONS

    @staticmethod
    def allowed_formats_message() -> str:
        exts = ", ".join(sorted(ALLOWED_EXTENSIONS))
        return f"الصيغ المدعومة: {exts}"

    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """
        استخراج النص الصافي — Raw Parsing فقط، بلا فحص دلالي أو LLM.
        يُرجع: (النص, نوع المصدر)
        """
        ext = DocumentProcessor.get_extension(filename)
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError(DocumentProcessor.allowed_formats_message())

        extractors: dict[str, Callable[[bytes], str]] = {
            ".docx": DocumentProcessor.extract_text_from_docx,
            ".pdf": DocumentProcessor.extract_text_from_pdf,
            ".txt": DocumentProcessor.extract_text_from_plain,
            ".md": DocumentProcessor.extract_text_from_plain,
            ".markdown": DocumentProcessor.extract_text_from_plain,
            ".json": DocumentProcessor.extract_text_from_json,
            ".rtf": DocumentProcessor.extract_text_from_rtf,
        }
        text = extractors[ext](file_bytes)
        if not text or not text.strip():
            raise ValueError("المستند فارغ أو لا يحتوي على نص قابل للاستخراج.")
        return text.strip(), EXTENSION_LABELS.get(ext, ext)

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """قراءة DOCX — فقرات وجداول (python-docx، تحميل كسول)."""
        from docx import Document

        try:
            doc = Document(io.BytesIO(file_bytes))
            full_text: list[str] = []

            for para in doc.paragraphs:
                line = para.text.strip()
                if line:
                    full_text.append(line)

            for table in doc.tables:
                for row in table.rows:
                    row_text: list[str] = []
                    for cell in row.cells:
                        val = cell.text.strip()
                        if val and (not row_text or row_text[-1] != val):
                            row_text.append(val)
                    if row_text:
                        full_text.append(" | ".join(row_text))

            extracted = "\n\n".join(full_text).strip()
            if not extracted:
                raise ValueError("المستند فارغ ولا يحتوي على نصوص في الفقرات أو الجداول.")
            return extracted
        except ValueError:
            raise
        except Exception as e:
            raise RuntimeError(f"فشل معالجة مستند DOCX: {str(e)}") from e

    @staticmethod
    def _extract_pdf_with_pypdf(file_bytes: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception:
                pass

        pages: list[str] = []
        for page in reader.pages:
            try:
                page_text = (page.extract_text() or "").strip()
            except Exception:
                page_text = ""
            if page_text:
                pages.append(page_text)
        return "\n\n".join(pages)

    @staticmethod
    def _extract_pdf_page_html(page) -> str:
        """استخراج عبر HTML ثم إزالة الوسوم — احتياطي لخطوط InDesign."""
        try:
            html = page.get_text("html") or ""
        except Exception:
            return ""
        if not html.strip():
            return ""
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"&nbsp;|&#\d+;", " ", text)
        text = re.sub(r"[ \t]+", " ", text)
        return "\n".join(ln.strip() for ln in text.splitlines() if ln.strip())

    @staticmethod
    def _extract_pdf_page_dict(page) -> str:
        """ترتيب بنية InDesign: خطوط → spans كما في PDF."""
        try:
            data = page.get_text("dict") or {}
        except Exception:
            return ""
        lines: list[str] = []
        for block in data.get("blocks") or []:
            if block.get("type") != 0:
                continue
            for line in block.get("lines") or []:
                spans = line.get("spans") or []
                chunk = "".join(s.get("text", "") for s in spans)
                chunk = chunk.strip()
                if chunk:
                    lines.append(chunk)
        return "\n".join(lines)

    @staticmethod
    def _extract_pdf_page_words(page) -> str:
        """ترتيب كلمات PDF الداخلي (block → line → word) — الأنسب لـ InDesign."""
        words = page.get_text("words") or []
        if not words:
            return ""
        rows: dict[tuple[int, int], list[tuple[int, str]]] = {}
        for w in words:
            if len(w) < 8:
                continue
            block_no, line_no, word_no = int(w[5]), int(w[6]), int(w[7])
            token = (w[4] or "").strip()
            if token:
                rows.setdefault((block_no, line_no), []).append((word_no, token))
        lines: list[str] = []
        for key in sorted(rows.keys()):
            tokens = sorted(rows[key], key=lambda t: t[0])
            lines.append(" ".join(t[1] for t in tokens))
        return "\n".join(lines)

    @staticmethod
    def _extract_pdf_page_blocks(page) -> str:
        """ترتيب كتل الصفحة حسب الموضع (احتياطي لـ RTL)."""
        blocks = page.get_text("blocks") or []
        rows: dict[int, list[tuple[float, str]]] = {}
        line_tol = 10

        for block in blocks:
            if len(block) < 7 or block[6] != 0:
                continue
            txt = (block[4] or "").strip()
            if not txt:
                continue
            y_key = int(block[1] / line_tol)
            rows.setdefault(y_key, []).append((float(block[0]), txt))

        if not rows:
            return (page.get_text("text") or "").strip()

        lines: list[str] = []
        for y_key in sorted(rows.keys()):
            parts = sorted(rows[y_key], key=lambda item: -item[0])
            lines.append(" ".join(part[1] for part in parts))
        return "\n".join(lines)

    @staticmethod
    def _extract_pdf_with_pymupdf(file_bytes: bytes, *, mode: str = "text") -> str:
        import fitz  # PyMuPDF

        pages: list[str] = []
        flags = getattr(fitz, "TEXT_DEHYPHENATE", 0)

        extractors = {
            "dict": DocumentProcessor._extract_pdf_page_dict,
            "words": DocumentProcessor._extract_pdf_page_words,
            "blocks": DocumentProcessor._extract_pdf_page_blocks,
            "html": DocumentProcessor._extract_pdf_page_html,
        }

        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                if mode in extractors:
                    page_text = extractors[mode](page)
                else:
                    try:
                        page_text = (page.get_text("text", flags=flags) or "").strip()
                    except TypeError:
                        page_text = (page.get_text("text") or "").strip()
                if page_text:
                    pages.append(page_text)
        return "\n\n".join(pages)

    @staticmethod
    def _extract_pdf_pymupdf_once(file_bytes: bytes, mode: str, doc) -> str:
        """استخراج كل الصفحات بنمط واحد من مستند مفتوح مسبقاً."""
        extractors = {
            "dict": DocumentProcessor._extract_pdf_page_dict,
            "words": DocumentProcessor._extract_pdf_page_words,
            "blocks": DocumentProcessor._extract_pdf_page_blocks,
            "html": DocumentProcessor._extract_pdf_page_html,
        }
        flags = 0
        try:
            import fitz

            flags = getattr(fitz, "TEXT_DEHYPHENATE", 0)
        except ImportError:
            pass

        pages: list[str] = []
        for page in doc:
            if mode in extractors:
                page_text = extractors[mode](page)
            else:
                try:
                    page_text = (page.get_text("text", flags=flags) or "").strip()
                except TypeError:
                    page_text = (page.get_text("text") or "").strip()
            if page_text:
                pages.append(page_text)
        return "\n\n".join(pages)

    @staticmethod
    def _open_pdf_document(file_bytes: bytes):
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if getattr(doc, "needs_pass", False):
            if not doc.authenticate(""):
                doc.close()
                raise ValueError(
                    "ملف PDF محمي بكلمة مرور. أزل الحماية أو صدّر نسخة Word (.docx)."
                )
        return doc

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """
        استخراج PDF: عدة طرق (InDesign) → أفضل نص خام → تنظيف مع fallback.
        """
        if not file_bytes or len(file_bytes) < 5:
            raise ValueError("ملف PDF فارغ أو تالف.")

        if not file_bytes[:5].startswith(b"%PDF"):
            raise ValueError("الملف ليس بصيغة PDF صالحة.")

        try:
            from pypdf import PdfReader  # noqa: F401 — فحص التثبيت
        except ImportError as e:
            raise RuntimeError(
                "مكتبة pypdf غير مثبتة على الخادم. من مجلد backend شغّل: "
                "venv\\Scripts\\pip install -r requirements.txt"
            ) from e

        errors: list[str] = []
        candidates: List[str] = []
        mode_stats: list[str] = []
        page_count = 0

        try:
            doc = DocumentProcessor._open_pdf_document(file_bytes)
            try:
                page_count = len(doc)
                for mode in ("dict", "words", "text", "blocks", "html"):
                    try:
                        chunk = DocumentProcessor._extract_pdf_pymupdf_once(
                            file_bytes, mode, doc
                        )
                    except Exception as e:
                        errors.append(f"{mode}: {e}")
                        continue
                    ar = arabic_char_count(chunk)
                    mode_stats.append(f"{mode}={ar}")
                    if chunk.strip():
                        candidates.append(chunk)
            finally:
                doc.close()
        except ValueError:
            raise
        except ImportError:
            errors.append("PyMuPDF غير مثبت")
        except Exception as e:
            errors.append(f"pymupdf: {e}")

        try:
            pypdf_text = DocumentProcessor._extract_pdf_with_pypdf(file_bytes)
            ar = arabic_char_count(pypdf_text)
            mode_stats.append(f"pypdf={ar}")
            if pypdf_text.strip():
                candidates.append(pypdf_text)
        except Exception as e:
            errors.append(f"pypdf: {e}")

        raw_best = pick_best_raw_pdf_text(candidates)
        if raw_best.strip():
            result = clean_pdf_arabic_text_with_fallback(raw_best)
            if result.strip() and (
                arabic_char_count(result) >= 15 or len(result.strip()) >= 80
            ):
                logger.info(
                    "pdf extract ok: %d pages, modes [%s], %d ar chars, %d total chars",
                    page_count,
                    ", ".join(mode_stats),
                    arabic_char_count(result),
                    len(result),
                )
                return result
            logger.warning(
                "pdf extract cleaned empty: pages=%d modes=[%s] raw_ar=%d",
                page_count,
                ", ".join(mode_stats),
                arabic_char_count(raw_best),
            )

        detail = (
            "تعذر استخراج نص من PDF. "
        )
        if page_count == 0:
            detail += "الملف لا يحتوي صفحات قابلة للقراءة. "
        elif max((arabic_char_count(c) for c in candidates), default=0) == 0:
            detail += (
                "لم يُعثر على حروف عربية — غالباً الملف ممسوحاً ضوئياً (صورة) "
                "أو خطوط مضمّنة غير قياسية. "
            )
        else:
            detail += "فشل التنظيف بعد الاستخراج الجزئي. "
        detail += "جرّب Word (.docx) من InDesign أو صدّر نص Unicode (.txt)."
        if mode_stats:
            detail += f" [تشخيص: {page_count} صفحة؛ {'; '.join(mode_stats[:6])}]"
        if errors:
            detail += f" ({'; '.join(errors[:2])})"
        raise ValueError(detail)

    @staticmethod
    def extract_text_from_plain(file_bytes: bytes) -> str:
        """TXT / Markdown — ترميز واحد سريع ثم fallback."""
        try:
            return file_bytes.decode("utf-8-sig")
        except UnicodeDecodeError:
            pass
        for encoding in ("utf-8", "cp1256", "windows-1256", "iso-8859-6", "latin-1"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    @staticmethod
    def extract_text_from_json(file_bytes: bytes) -> str:
        """JSON — قراءة ميكانيكية فورية، بلا json_document_parser."""
        try:
            payload = _loads_json(file_bytes)
        except (json.JSONDecodeError, ValueError):
            return DocumentProcessor.extract_text_from_plain(file_bytes)

        if isinstance(payload, dict):
            joined = _json_sections_to_text(payload)
            if joined:
                return joined
            text_field = payload.get("text")
            if isinstance(text_field, str) and text_field.strip():
                return text_field
            # آخر ملاذ: نص خام بدون pretty-print (كان يبطّئ الملفات الضخمة)
            if orjson is not None:
                return orjson.dumps(payload).decode("utf-8")
            return json.dumps(payload, ensure_ascii=False, separators=(",", ":"))

        if isinstance(payload, list):
            if orjson is not None:
                return orjson.dumps(payload).decode("utf-8")
            return json.dumps(payload, ensure_ascii=False, separators=(",", ":"))

        return str(payload)

    @staticmethod
    def extract_text_from_rtf(file_bytes: bytes) -> str:
        """RTF — عبر striprtf أو تنظيف بسيط."""
        try:
            from striprtf.striprtf import rtf_to_text

            raw = file_bytes.decode("cp1256", errors="replace")
            return rtf_to_text(raw)
        except ImportError:
            raw = file_bytes.decode("utf-8", errors="replace")
            text = re.sub(r"\\[a-z]+\d*\s?", "", raw)
            text = re.sub(r"[{}]", "", text)
            text = re.sub(r"\s+", " ", text)
            return text.strip()
