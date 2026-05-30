# -*- coding: utf-8 -*-

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


ARABIC_FONT = "Traditional Arabic"
FALLBACK_FONT = "Arial"


def _set_paragraph_rtl(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_run_font(run, size_pt: int = 14, bold: bool = False, color: Optional[RGBColor] = None) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = FALLBACK_FONT
    run.font.rtl = True
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FALLBACK_FONT)
    r_fonts.set(qn("w:hAnsi"), FALLBACK_FONT)
    r_fonts.set(qn("w:cs"), ARABIC_FONT)
    r_fonts.set(qn("w:rtl"), ARABIC_FONT)
    r_pr.insert(0, r_fonts)
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    run = p.add_run(text)
    sizes = {1: 22, 2: 16, 3: 14}
    colors = {
        1: RGBColor(0xB4, 0x53, 0x09),
        2: RGBColor(0x1E, 0x3A, 0x5F),
        3: RGBColor(0x33, 0x41, 0x55),
    }
    _set_run_font(run, size_pt=sizes.get(level, 14), bold=True, color=colors.get(level))
    p.paragraph_format.space_before = Pt(14 if level > 1 else 6)
    p.paragraph_format.space_after = Pt(8)


def _add_body(doc: Document, text: str, size: int = 14, italic: bool = False) -> None:
    if not text or not str(text).strip():
        return
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    run = p.add_run(str(text).strip())
    _set_run_font(run, size_pt=size, bold=False)
    run.italic = italic
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)


def _add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    label_run = p.add_run(f"{label}: ")
    _set_run_font(label_run, size_pt=12, bold=True, color=RGBColor(0x64, 0x74, 0x8B))
    value_run = p.add_run(value)
    _set_run_font(value_run, size_pt=13)
    p.paragraph_format.space_after = Pt(4)


class ConsolidationDocxExporter:
    """تصدير نتائج الصهر الديناميكي إلى مستند Word منسّق (RTL)."""

    @staticmethod
    def build(
        discovered_structure: Dict[str, Any],
        title: str = "جوهر المخطوطة — مدونة الخليل",
        source_filename: Optional[str] = None,
    ) -> bytes:
        doc = Document()

        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.left_margin = Cm(2.0)

        _add_heading(doc, title, level=1)
        _add_body(
            doc,
            "مستخلص سيادي — الصهر الديناميكي وعكس الهندسة الدلالية v4.0",
            size=12,
            italic=True,
        )

        meta = discovered_structure.get("_metadata") or {}
        meta_bits = [
            f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        ]
        if source_filename:
            meta_bits.append(f"المصدر: {source_filename}")
        if meta.get("clustering_mode"):
            meta_bits.append(f"وضع التجميع: {meta['clustering_mode']}")
        if meta.get("engine_description") or meta.get("engine_utilized"):
            meta_bits.append(f"المحرك: {meta.get('engine_description') or meta.get('engine_utilized')}")
        _add_body(doc, " | ".join(meta_bits), size=10, italic=True)

        doc.add_paragraph()

        ideas: List[Dict] = discovered_structure.get("core_ideas") or []
        _add_heading(doc, "البطاقات المعرفية السيادية", level=2)

        if not ideas:
            _add_body(doc, "لم تُستخلص بطاقات معرفية.", italic=True)
        else:
            for idea in ideas:
                section_title = idea.get("section_title") or f"محور {idea.get('id', '')}"
                _add_heading(doc, f"{idea.get('id', '')}. {section_title}", level=3)

                sovereign = idea.get("sovereign_idea") or idea.get("idea") or ""
                _add_label_value(doc, "الصياغة السيادية", sovereign)

                layers = idea.get("layers") or {}
                if layers.get("conceptual_framework"):
                    _add_label_value(doc, "الإطار المفاهيمي", layers["conceptual_framework"])
                if layers.get("practical_applications"):
                    _add_label_value(doc, "التطبيقات العملية", layers["practical_applications"])

                styles = idea.get("discovered_styles") or []
                if styles:
                    _add_label_value(doc, "الأنماط المكتشفة", " · ".join(styles))

                doc.add_paragraph()

        ledger = discovered_structure.get("numerical_ledger") or []
        _add_heading(doc, "الكشاف الرقمي والتواريخ", level=2)

        if ledger:
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "القيمة"
            hdr[1].text = "السياق"
            for cell in hdr:
                for p in cell.paragraphs:
                    _set_paragraph_rtl(p)
                    for run in p.runs:
                        _set_run_font(run, size_pt=11, bold=True)

            for item in ledger[:20]:
                row = table.add_row().cells
                row[0].text = str(item.get("value", ""))
                row[1].text = str(item.get("context", ""))
                for cell in row:
                    for p in cell.paragraphs:
                        _set_paragraph_rtl(p)
                        for run in p.runs:
                            _set_run_font(run, size_pt=11)
        else:
            _add_body(doc, "لا توجد أرقام أو تواريخ بارزة.", italic=True)

        keywords = discovered_structure.get("sovereign_keywords") or []
        _add_heading(doc, "الكلمات المفتاحية السيادية", level=2)
        if keywords:
            _add_body(doc, "، ".join(str(k) for k in keywords))
        else:
            _add_body(doc, "—", italic=True)

        styles_found = discovered_structure.get("discovered_styles") or []
        if styles_found:
            _add_heading(doc, "الأنماط الأسلوبية العامة", level=2)
            _add_body(doc, "، ".join(styles_found))

        doc.add_paragraph()
        token_usage = meta.get("token_usage") or {}
        total_tkn = token_usage.get("total_tokens") or meta.get("tokens_consumed", 0)
        est = " (تقدير)" if token_usage.get("estimated") else ""
        token_line = f"التوكنات: {total_tkn}{est}"
        if token_usage.get("input_tokens") is not None:
            token_line += (
                f" | مدخل: {token_usage.get('input_tokens', 0)}"
                f" | مخرج: {token_usage.get('output_tokens', 0)}"
                f" | استدعاءات LLM: {token_usage.get('llm_calls', '—')}"
            )
        audit = "ناجح" if meta.get("audit_passed") else "تحذيرات — يُراجع يدوياً"
        _add_body(doc, token_line, size=10)
        _add_body(
            doc,
            f"مدونة الخليل للتحرير اللغوي — Copyright © {datetime.now().year} | التدقيق: {audit}",
            size=9,
            italic=True,
        )

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
